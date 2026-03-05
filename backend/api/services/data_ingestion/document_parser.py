from pathlib import Path


def parse_document(file_path: str) -> str:
    """Parse PDF, DOCX, or TXT file and return plain text."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == '.pdf':
        return _parse_pdf(file_path)
    elif ext == '.docx':
        return _parse_docx(file_path)
    elif ext == '.txt':
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported document type: {ext}")


def _parse_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return '\n'.join(text_parts).strip()


def _parse_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = '\t'.join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return '\n'.join(parts)


def _parse_txt(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read().strip()
